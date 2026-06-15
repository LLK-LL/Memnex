from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

CJK_SAMPLE = "\u4e2d\u6587\u5185\u5bb9 ABC"
SIMSUN = "\u5b8b\u4f53"


def make_docx_missing_east_asia(path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(CJK_SAMPLE)
    run.font.name = "Times New Roman"
    doc.save(path)


def make_docx_with_east_asia(path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), SIMSUN)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(CJK_SAMPLE)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), SIMSUN)
    doc.save(path)


if __name__ == "__main__":
    make_docx_missing_east_asia(Path("sample-missing.docx"))
    make_docx_with_east_asia(Path("sample-ok.docx"))
